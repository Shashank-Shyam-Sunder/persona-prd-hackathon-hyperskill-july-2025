# src/prd_generator.py

import os
import pandas as pd
from typing import List
from dotenv import load_dotenv
from docx import Document

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from persona_config import PERSONA_TO_FOLDER, PERSONA_DISPLAY_NAMES

# ───────────────────────────────────────────────
# 🔐 Load environment variables
# ───────────────────────────────────────────────
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables.")

# ───────────────────────────────────────────────
# 🧠 Gemini model initialization
# ───────────────────────────────────────────────
model = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash",
    google_api_key=GOOGLE_API_KEY
)

# ───────────────────────────────────────────────
# 📄 Prompt template
# ───────────────────────────────────────────────
PRD_PROMPT_TEMPLATE = """
You are a senior Product Manager AI assistant at PersonaPRD, an AI-powered product discovery platform.

Your task is to generate a structured Product Requirements Document (PRD) draft based on the following pain point summaries collected from user community data.

**Persona:** {persona_name}

**Pain Points:**
{pain_points}

**Instructions:**

Write a PRD draft that includes:

1. **Problem Summary:** A clear summary of the problem(s) these pain points represent. Use simple, direct language.  
2. **Why This Problem Matters:** Explain why this problem is significant specifically for the **{persona_name}** persona. Highlight the impact on productivity, workflows, or business goals.  
3. **Potential Solution Overview:** Provide a concise solution concept that addresses these pain points.  
4. **Suggested MVP Features:** List 3–5 minimum viable product features as bullet points, phrased as actionable features (e.g. “Feature X does Y to solve Z”).  
5. **Next Steps:** Outline immediate steps the team should take to validate and build this solution (e.g. user interviews, prototype, sprint planning).

**Tone guidelines:**
- Use clear, professional, and confident language.  
- Avoid generic filler phrases.  
- Tailor content to a product team preparing for sprint planning.

PRD Draft:
"""

prompt = PromptTemplate(
    input_variables=["pain_points", "persona_name"],
    template=PRD_PROMPT_TEMPLATE,
)

prd_chain = prompt | model

# ───────────────────────────────────────────────
# 🔧 Core Functions
# ───────────────────────────────────────────────

def generate_prd(pain_point_summaries: List[str], persona_display_name: str) -> str:
    combined = "\n".join([f"- {pp}" for pp in pain_point_summaries])
    response = prd_chain.invoke({"pain_points": combined, "persona_name": persona_display_name})
    return response.content.strip()


def load_summaries(persona_key: str, subreddit_file: str) -> pd.DataFrame:
    persona_folder = PERSONA_TO_FOLDER.get(persona_key)
    if not persona_folder:
        raise ValueError(f"Unknown persona key: {persona_key}")

    subreddit_folder = subreddit_file.replace(".json", "")
    summaries_path = os.path.join("data", "processed", persona_folder, subreddit_folder, "pain_point_summaries.csv")
    if not os.path.exists(summaries_path):
        raise FileNotFoundError(f"Pain point summaries not found at {summaries_path}")
    return pd.read_csv(summaries_path)


def select_pain_points(df: pd.DataFrame) -> List[str]:
    print("\n📌 Found", len(df), "pain point summaries. Here they are:\n")
    for idx, row in df.iterrows():
        print(f"[{row['cluster_id']}] {row['pain_point_summary']}\n")

    selected = input("👉 Enter comma-separated cluster IDs to include in PRD (e.g. 0,2,5): ")
    cluster_ids = [int(cid.strip()) for cid in selected.split(",") if cid.strip().isdigit()]
    return df[df["cluster_id"].isin(cluster_ids)]["pain_point_summary"].tolist()


def save_prd_to_docx(prd_text: str, output_path: str):
    doc = Document()
    doc.add_heading("Product Requirements Document (Generated)", level=1)
    for line in prd_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(output_path)
    print(f"\n✅ PRD saved to: {output_path}")


def run_prd_generator(persona_key: str, subreddit_file: str, persona_display_name: str):
    print(f"\n📝 Generating PRD for Persona: {persona_display_name} / Subreddit: {subreddit_file}\n")

    df = load_summaries(persona_key, subreddit_file)
    selected_summaries = select_pain_points(df)
    prd_text = generate_prd(selected_summaries, persona_display_name)

    print("\n=== Generated PRD Draft ===\n")
    print(prd_text)

    persona_folder = PERSONA_TO_FOLDER[persona_key]
    subreddit_folder = subreddit_file.replace(".json", "")
    output_folder = os.path.join("data", "processed", persona_folder, subreddit_folder)
    os.makedirs(output_folder, exist_ok=True)

    save_path = os.path.join(output_folder, "PRD_DRAFT.docx")
    save_prd_to_docx(prd_text, save_path)


# ───────────────────────────────────────────────
# 🧪 Manual Testing (Optional)
# ───────────────────────────────────────────────
if __name__ == "__main__":
    persona_key = "vibecoding"
    subreddit_file = "reddit_PromptEngineering_hot_500.json"
    display_name = PERSONA_DISPLAY_NAMES[persona_key]
    run_prd_generator(persona_key, subreddit_file, display_name)
